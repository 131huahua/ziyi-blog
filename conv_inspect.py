# 查看 docx 结构样例
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document

doc = Document(r'C:\Users\zzy\Desktop\50香肠配方\哈尔滨红肠原料配方.docx')
for i, p in enumerate(doc.paragraphs[:40]):
    t = p.text.strip()
    if t:
        print(f'[{i}] style={p.style.name!r}: {t[:60]}')
print('--- tables:', len(doc.tables))
for tb in doc.tables[:2]:
    for row in tb.rows[:5]:
        print(' | '.join(c.text.strip()[:20] for c in row.cells))
